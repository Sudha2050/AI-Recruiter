import docx
import sys
from pathlib import Path

def convert_docx_to_md(docx_path: Path, md_path: Path):
    doc = docx.Document(docx_path)
    lines = []
    for p in doc.paragraphs:
        # Simple conversions for headings
        if p.style.name.startswith('Heading 1'):
            lines.append(f"\n# {p.text}\n")
        elif p.style.name.startswith('Heading 2'):
            lines.append(f"\n## {p.text}\n")
        elif p.style.name.startswith('Heading 3'):
            lines.append(f"\n### {p.text}\n")
        else:
            lines.append(p.text)
    
    # Process tables if any
    for table in doc.tables:
        lines.append("\n| " + " | ".join(cell.text.replace('\n', ' ') for cell in table.rows[0].cells) + " |")
        lines.append("| " + " | ".join("---" for _ in table.rows[0].cells) + " |")
        for row in table.rows[1:]:
            lines.append("| " + " | ".join(cell.text.replace('\n', ' ') for cell in row.cells) + " |")
        lines.append("\n")

    md_path.write_text('\n'.join(lines), encoding='utf-8')
    print(f"Converted {docx_path.name} to {md_path.name}")

if __name__ == '__main__':
    artifact_dir = Path(r"C:\Users\gnoob\AppData\Local\Temp" if not sys.argv[1:] else sys.argv[1])
    dataset_dir = Path("data/India_runs_data_and_ai_challenge")
    
    convert_docx_to_md(dataset_dir / "README.docx", artifact_dir / "README_doc.md")
    convert_docx_to_md(dataset_dir / "redrob_signals_doc.docx", artifact_dir / "redrob_signals_doc.md")
    convert_docx_to_md(dataset_dir / "submission_spec.docx", artifact_dir / "submission_spec_doc.md")
